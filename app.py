#!/usr/bin/env python3
"""
Aplicación Streamlit para transcribir archivos de audio usando OpenAI Whisper.
Permite subir audios, elegir modelo/idioma, transcribir y descargar en TXT/DOCX.
"""

import os
import tempfile
import streamlit as st
import whisper
from docx import Document
from io import BytesIO
import time
import subprocess
import sys
import re
from io import StringIO

def export_docx(text: str, segments=None) -> bytes:
    """
    Exporta la transcripción a formato DOCX con timestamps opcionales.
    
    Args:
        text (str): Texto de la transcripción
        segments (list): Lista de segmentos con timestamps (opcional)
    
    Returns:
        bytes: Contenido del archivo DOCX
    """
    doc = Document()
    doc.add_heading('Transcripción de Audio', 0)
    
    # Añadir información general
    doc.add_heading('Información', level=1)
    doc.add_paragraph(f'Fecha de transcripción: {time.strftime("%Y-%m-%d %H:%M:%S")}')
    if segments:
        doc.add_paragraph(f'Número de segmentos: {len(segments)}')
    
    # Añadir transcripción completa
    doc.add_heading('Transcripción Completa', level=1)
    doc.add_paragraph(text)
    
    # Añadir segmentos con timestamps si están disponibles
    if segments:
        doc.add_heading('Segmentos con Timestamps', level=1)
        for i, segmento in enumerate(segments):
            inicio = int(segmento['start'])
            fin = int(segmento['end'])
            texto_seg = segmento['text']
            
            # Formatear tiempo
            inicio_min = inicio // 60
            inicio_sec = inicio % 60
            fin_min = fin // 60
            fin_sec = fin % 60
            
            doc.add_paragraph(f'[{inicio_min:02d}:{inicio_sec:02d} - {fin_min:02d}:{fin_sec:02d}] {texto_seg}')
    
    # Guardar en BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def format_time(seconds):
    """Formatea segundos a formato MM:SS"""
    minutes = int(seconds // 60)
    seconds = int(seconds % 60)
    return f"{minutes:02d}:{seconds:02d}"

def setup_ffmpeg():
    """Configura ffmpeg para que esté disponible en el PATH"""
    try:
        # Verificar si ffmpeg está disponible
        result = subprocess.run(['ffmpeg', '-version'], 
                             capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            return True
    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass
    
    # Si no está disponible, intentar usar los ejecutables locales
    try:
        local_ffmpeg = os.path.join(os.getcwd(), 'ffmpeg', 'ffmpeg.exe')
        if os.path.exists(local_ffmpeg):
            # Añadir al PATH temporalmente
            ffmpeg_dir = os.path.join(os.getcwd(), 'ffmpeg')
            os.environ['PATH'] = ffmpeg_dir + os.pathsep + os.environ['PATH']
            
            # Verificar que funciona
            result = subprocess.run([local_ffmpeg, '-version'], 
                                 capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                return True
    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass
    
    return False

# Configuración de la página
st.set_page_config(
    page_title='Transcriptor Whisper',
    page_icon='🎙️',
    layout='wide',
    initial_sidebar_state='expanded'
)

# Verificar y configurar ffmpeg
@st.cache_resource
def init_ffmpeg():
    """Inicializa ffmpeg y retorna si está disponible"""
    return setup_ffmpeg()

# Verificar ffmpeg al cargar la página
ffmpeg_available = init_ffmpeg()
if not ffmpeg_available:
    st.warning("⚠️ FFmpeg no está disponible. La transcripción puede fallar.")

# Cache del modelo para evitar recargas
@st.cache_resource
def load_whisper_model(model_name):
    """Carga el modelo de Whisper y lo mantiene en caché"""
    try:
        return whisper.load_model(model_name)
    except Exception as e:
        st.error(f"Error al cargar el modelo: {e}")
        return None

# Título principal
st.title('🎙️ Transcriptor de Audio con Whisper')
st.markdown('---')

# Inicializar estado de procesamiento
if 'procesando' not in st.session_state:
    st.session_state.procesando = False

# Sidebar con opciones
st.sidebar.header('⚙️ Configuración')

# Mostrar advertencia si está procesando
if st.session_state.procesando:
    st.sidebar.warning('⚠️ Transcripción en proceso... Por favor espera.')

# Selector de modelo
modelo_options = ['tiny', 'base', '🚫 small (no disponible en CPU gratuita)']
modelo = st.sidebar.selectbox(
    '🤖 Modelo de Whisper',
    modelo_options,
    index=1,
    help='Modelos más grandes = mejor precisión pero más lento',
    disabled=st.session_state.procesando
)

# Convertir la selección a nombre real del modelo
if modelo.startswith('🚫'):
    st.sidebar.error('⚠️ El modelo "small" requiere demasiados recursos y no funciona bien en CPU gratuita. Por favor, selecciona "tiny" o "base".')
    modelo_real = 'base'  # Fallback
    modelo_disabled = True
else:
    modelo_real = modelo
    modelo_disabled = False

# Selector de idioma
idioma = st.sidebar.selectbox(
    '🌍 Idioma',
    ['auto', 'es', 'en', 'fr', 'pt', 'it', 'de'],
    index=1,
    help='Dejar en "auto" para detección automática',
    disabled=st.session_state.procesando
)

# Información sobre modelos
st.sidebar.markdown('### 📊 Información de Modelos')
model_info = {
    'tiny': '~39 MB - Muy rápido, precisión básica',
    'base': '~74 MB - Equilibrio velocidad/precisión (recomendado)',
}
if modelo_real in model_info:
    st.sidebar.info(f"**{modelo_real}**: {model_info[modelo_real]}")

# Advertencia de rendimiento en CPU
st.sidebar.markdown('### ⚡ Rendimiento en CPU')
tiempo_estimado = {
    'tiny': '3-8 min',
    'base': '10-25 min',
}
if modelo_real in tiempo_estimado:
    st.sidebar.warning(f"⏱️ Tiempo estimado para audio de 200MB: **{tiempo_estimado[modelo_real]}**")
st.sidebar.info("💡 **Consejo**: Usa archivos más pequeños o modelo 'tiny' para pruebas rápidas.")

# Área principal
col1, col2 = st.columns([2, 1])

with col1:
    st.header('📁 Subir Archivo de Audio')
    
    # File uploader
    archivo = st.file_uploader(
        'Selecciona un archivo de audio',
        type=['mp3', 'wav', 'm4a', 'flac', 'ogg', 'mp4', 'webm', 'mkv'],
        help='Recomendado: < 50MB para mejor rendimiento en CPU. Máximo: 200MB',
        disabled=st.session_state.procesando
    )
    
    if archivo:
        tamano_mb = archivo.size / (1024*1024)
        st.success(f'✅ Archivo cargado: {archivo.name}')
        
        # Advertencia por tamaño
        if tamano_mb > 100:
            st.warning(f'⚠️ Tamaño: {tamano_mb:.1f} MB - La transcripción puede tardar mucho en CPU')
        elif tamano_mb > 50:
            st.info(f'📏 Tamaño: {tamano_mb:.1f} MB - Tiempo de transcripción moderado')
        else:
            st.info(f'📏 Tamaño: {tamano_mb:.1f} MB - Tamaño óptimo para CPU')
        
        # Botón de transcripción
        col_btn1, col_btn2 = st.columns([3, 1])
        with col_btn1:
            iniciar_transcripcion = st.button('🚀 Transcribir Audio', type='primary', use_container_width=True, disabled=modelo_disabled or st.session_state.procesando)
        
        if iniciar_transcripcion:
            if archivo.size > 200 * 1024 * 1024:  # 200MB
                st.error('❌ El archivo es demasiado grande. Máximo 200MB.')
            elif not ffmpeg_available:
                st.error('❌ FFmpeg no está disponible. No se puede procesar el audio.')
                st.info('💡 Intenta recargar la página o contacta al administrador.')
            elif modelo_disabled:
                st.error('❌ El modelo seleccionado no está disponible. Por favor, elige "tiny" o "base".')
            else:
                # Activar estado de procesamiento
                st.session_state.procesando = True
                
                # Crear archivo temporal
                with tempfile.NamedTemporaryFile(delete=False, suffix=f'_{archivo.name}') as tmp:
                    tmp.write(archivo.read())
                    ruta_temp = tmp.name
                
                try:
                    # Contenedores para progreso y logs
                    st.markdown('---')
                    st.subheader('🔄 Proceso de Transcripción')
                    
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    # Contenedor simple para logs de Whisper
                    whisper_logs_container = st.empty()
                    
                    # Calcular tiempo estimado
                    duracion_estimada_min = {
                        'tiny': tamano_mb * 0.5,  # ~30 seg por MB
                        'base': tamano_mb * 1.2,  # ~1.2 min por MB
                    }
                    tiempo_est = duracion_estimada_min.get(modelo_real, tamano_mb * 1.0)
                    
                    # Advertencia sobre cancelación
                    st.warning('⚠️ **Importante**: Una vez iniciado, el proceso no se puede cancelar. Si necesitas detenerlo, recarga la página (perderás el progreso).')
                    
                    # Info inicial
                    st.info(f'📊 **Modelo**: {modelo_real} | **Tamaño**: {tamano_mb:.1f} MB | **Tiempo estimado**: ~{tiempo_est:.1f} min')
                    
                    # Cargar modelo (cacheado)
                    status_text.info('🔄 Paso 1/3: Cargando modelo Whisper...')
                    progress_bar.progress(10)
                    
                    tiempo_inicio = time.time()
                    model = load_whisper_model(modelo_real)
                    if model is None:
                        st.error('❌ No se pudo cargar el modelo.')
                        st.stop()
                    
                    progress_bar.progress(30)
                    status_text.info('🎵 Paso 2/3: Transcribiendo audio... (Esto puede tardar varios minutos)')
                    
                    # Configurar parámetros de transcripción optimizados para CPU
                    # Capturar el progreso de Whisper
                    progreso_info = {'porcentaje': 0, 'ultimo_log': ''}
                    
                    # Redirigir stderr para capturar logs de Whisper
                    import contextlib
                    
                    class ProgressCapture:
                        def __init__(self):
                            self.logs = []
                            
                        def write(self, text):
                            if text.strip():
                                self.logs.append(text)
                                # Buscar porcentajes en los logs (formato: XX% o XX%)
                                match = re.search(r'(\d+)%', text)
                                if match:
                                    porcentaje = int(match.group(1))
                                    progreso_info['porcentaje'] = porcentaje
                                    progreso_info['ultimo_log'] = text.strip()
                        
                        def flush(self):
                            pass
                    
                    transcribe_kwargs = {
                        'fp16': False,  # Forzar FP32 en CPU
                        'verbose': True,  # Activar para capturar progreso
                    }
                    if idioma != 'auto':
                        transcribe_kwargs['language'] = idioma
                    
                    # Transcribir con captura de progreso real
                    tiempo_transcripcion_inicio = time.time()
                    
                    # Contenedores para progreso
                    transcription_progress = st.empty()
                    detail_progress = st.empty()
                    
                    # Ejecutar transcripción en un thread
                    import threading
                    resultado_container = {'resultado': None, 'error': None, 'completado': False}
                    progress_capture = ProgressCapture()
                    
                    def transcribe_thread():
                        try:
                            # Redirigir stderr para capturar logs de Whisper
                            old_stderr = sys.stderr
                            sys.stderr = progress_capture
                            
                            resultado_container['resultado'] = model.transcribe(ruta_temp, **transcribe_kwargs)
                            resultado_container['completado'] = True
                            
                            # Restaurar stderr
                            sys.stderr = old_stderr
                        except Exception as e:
                            sys.stderr = old_stderr
                            resultado_container['error'] = str(e)
                            resultado_container['completado'] = True
                    
                    # Iniciar transcripción en background
                    thread = threading.Thread(target=transcribe_thread)
                    thread.start()
                    
                    # Actualizar progreso basado en logs reales de Whisper
                    progreso_anterior = 0
                    logs_whisper_texto = []
                    
                    while not resultado_container['completado']:
                        tiempo_transcurrido = time.time() - tiempo_transcripcion_inicio
                        
                        # Obtener porcentaje real de Whisper
                        porcentaje_whisper = progreso_info['porcentaje']
                        ultimo_log = progreso_info['ultimo_log']
                        
                        if porcentaje_whisper > 0:
                            # Usar el progreso real de Whisper (mapear de 0-100% a 40-90%)
                            progreso_real = 40 + int((porcentaje_whisper / 100) * 50)
                            progress_bar.progress(progreso_real)
                            
                            transcription_progress.info(f'🎵 Transcribiendo... **{porcentaje_whisper}%** completado')
                            
                            # Calcular tiempo restante basado en el progreso real
                            if porcentaje_whisper > 5:  # Evitar divisiones por números muy pequeños
                                tiempo_por_porcentaje = tiempo_transcurrido / porcentaje_whisper
                                tiempo_restante_seg = tiempo_por_porcentaje * (100 - porcentaje_whisper)
                                minutos = int(tiempo_restante_seg // 60)
                                segundos = int(tiempo_restante_seg % 60)
                                detail_progress.success(f'⏱️ Tiempo transcurrido: {int(tiempo_transcurrido)}s | Tiempo estimado restante: **~{minutos}m {segundos}s**')
                            else:
                                detail_progress.info(f'⏱️ Tiempo transcurrido: {int(tiempo_transcurrido)}s | Calculando tiempo restante...')
                            
                            # Actualizar logs de Whisper en la cajita
                            if ultimo_log and (not logs_whisper_texto or logs_whisper_texto[-1] != ultimo_log):
                                logs_whisper_texto.append(ultimo_log)
                                # Mantener solo los últimos 10 logs para no sobrecargar
                                if len(logs_whisper_texto) > 10:
                                    logs_whisper_texto.pop(0)
                                
                                # Mostrar logs en cajita
                                with whisper_logs_container.container():
                                    st.markdown('##### 📋 Logs de Whisper')
                                    logs_text = '\n'.join([f'• {log}' for log in logs_whisper_texto[-5:]])  # Últimos 5
                                    st.code(logs_text, language=None)
                        else:
                            # Si aún no hay progreso, mostrar mensaje de inicio
                            progreso_estimado = min(90, 40 + int(tiempo_transcurrido * 2))
                            progress_bar.progress(progreso_estimado)
                            transcription_progress.info(f'🎵 Inicializando transcripción...')
                            detail_progress.info(f'⏱️ Tiempo transcurrido: {int(tiempo_transcurrido)}s')
                        
                        time.sleep(0.5)  # Actualizar cada medio segundo para más fluidez
                    
                    # Esperar a que termine el thread
                    thread.join()
                    
                    # Limpiar mensajes de progreso
                    transcription_progress.empty()
                    detail_progress.empty()
                    whisper_logs_container.empty()
                    
                    # Verificar errores
                    if resultado_container['error']:
                        st.session_state.procesando = False  # Desactivar estado
                        st.error(f'❌ Error durante la transcripción: {resultado_container["error"]}')
                        st.stop()
                    
                    resultado = resultado_container['resultado']
                    tiempo_transcripcion = time.time() - tiempo_transcripcion_inicio
                    progress_bar.progress(90)
                    
                    status_text.info('📝 Paso 3/3: Generando resultados...')
                    
                    progress_bar.progress(100)
                    status_text.success('✅ ¡Transcripción completada exitosamente!')
                    
                    tiempo_total = time.time() - tiempo_inicio
                    st.success(f'🎉 Proceso completado en {tiempo_total/60:.1f} minutos')
                    
                    # Desactivar estado de procesamiento
                    st.session_state.procesando = False
                    
                    # Mostrar resultados
                    st.success('🎉 ¡Transcripción completada exitosamente!')
                    
                    # Información del resultado
                    col_info1, col_info2, col_info3 = st.columns(3)
                    with col_info1:
                        st.metric('Idioma detectado', resultado.get('language', 'No detectado'))
                    with col_info2:
                        st.metric('Segmentos', len(resultado.get('segments', [])))
                    with col_info3:
                        duracion = resultado.get('segments', [{}])[-1].get('end', 0) if resultado.get('segments') else 0
                        st.metric('Duración', format_time(duracion))
                    
                    # Mostrar transcripción
                    st.header('📝 Resultado de la Transcripción')
                    texto_completo = resultado['text']
                    st.text_area(
                        'Transcripción completa',
                        texto_completo,
                        height=300,
                        help='Puedes copiar este texto directamente'
                    )
                    
                    # Mostrar segmentos con timestamps
                    if resultado.get('segments'):
                        st.header('⏱️ Segmentos con Timestamps')
                        
                        # Mostrar solo los primeros 10 segmentos para no sobrecargar
                        segmentos_mostrar = resultado['segments'][:10]
                        
                        for i, segmento in enumerate(segmentos_mostrar):
                            inicio = segmento['start']
                            fin = segmento['end']
                            texto_seg = segmento['text']
                            
                            st.markdown(f"**{format_time(inicio)} - {format_time(fin)}**: {texto_seg}")
                        
                        if len(resultado['segments']) > 10:
                            st.info(f'... y {len(resultado["segments"]) - 10} segmentos más')
                    
                    # Botones de descarga
                    st.header('💾 Descargar Resultados')
                    
                    col_download1, col_download2 = st.columns(2)
                    
                    with col_download1:
                        st.download_button(
                            '📄 Descargar TXT',
                            texto_completo,
                            file_name=f'transcripcion_{archivo.name.split(".")[0]}.txt',
                            mime='text/plain',
                            use_container_width=True
                        )
                    
                    with col_download2:
                        docx_content = export_docx(texto_completo, resultado.get('segments'))
                        st.download_button(
                            '📝 Descargar DOCX',
                            docx_content,
                            file_name=f'transcripcion_{archivo.name.split(".")[0]}.docx',
                            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            use_container_width=True
                        )
                    
                except Exception as e:
                    st.session_state.procesando = False  # Desactivar estado en caso de error
                    st.error(f'❌ Error durante la transcripción: {str(e)}')
                    st.info('💡 Sugerencias: Verifica que el archivo no esté corrupto y que sea un formato de audio válido.')
                
                finally:
                    # Limpiar archivo temporal
                    try:
                        os.unlink(ruta_temp)
                    except:
                        pass
                    # Asegurar que el estado se desactiva siempre
                    st.session_state.procesando = False

with col2:
    st.header('ℹ️ Información')
    
    st.markdown("""
    ### 🎯 Cómo usar:
    1. **Selecciona** el modelo y idioma en la barra lateral
    2. **Sube** tu archivo de audio
    3. **Haz clic** en "Transcribir Audio"
    4. **Espera** a que se complete (puede tardar varios minutos)
    5. **Descarga** el resultado en TXT o DOCX
    
    ### ⚡ Consejos de rendimiento:
    - **tiny**: Más rápido, menos preciso
    - **base**: Equilibrio perfecto (recomendado)
    - **small**: Más preciso, más lento
    
    ### 📁 Formatos soportados:
    MP3, WAV, M4A, FLAC, OGG, MP4, WEBM, MKV
    """)
    
    st.markdown('---')
    
    st.markdown("""
    ### 🚀 Despliegue
    Esta app está optimizada para:
    - **Streamlit Cloud** (CPU)
    - **Hugging Face Spaces** (GPU opcional)
    """)

# Footer
st.markdown('---')
st.markdown(
    'Desarrollado con ❤️ usando [Streamlit](https://streamlit.io) y [OpenAI Whisper](https://github.com/openai/whisper)'
)
