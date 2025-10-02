"""
Funciones para exportar transcripciones en diferentes formatos.
"""
from docx import Document
from io import BytesIO
import time


def export_txt(text, segments=None):
    """
    Exporta la transcripción a formato TXT.
    
    Args:
        text (str): Texto de la transcripción
        segments (list): Lista de segmentos con timestamps (opcional)
    
    Returns:
        str: Contenido del archivo TXT
    """
    output = []
    output.append("=== TRANSCRIPCIÓN ===\n\n")
    output.append(text)
    
    if segments:
        output.append("\n\n=== SEGMENTOS CON TIMESTAMPS ===\n\n")
        for segmento in segments:
            inicio = int(segmento['start'])
            fin = int(segmento['end'])
            texto_seg = segmento['text']
            
            inicio_min = inicio // 60
            inicio_sec = inicio % 60
            fin_min = fin // 60
            fin_sec = fin % 60
            
            output.append(f"[{inicio_min:02d}:{inicio_sec:02d} - {fin_min:02d}:{fin_sec:02d}] {texto_seg}\n")
    
    return ''.join(output)


def export_docx(text, segments=None):
    """
    Exporta la transcripción a formato DOCX con timestamps opcionales.
    
    Args:
        text (str): Texto de la transcripción
        segments (list): Lista de segmentos con timestamps (opcional)
    
    Returns:
        bytes: Contenido del archivo DOCX
    """
    doc = Document()
    doc.add_heading('Transcripción de Audio', 0)
    
    # Añadir información general
    doc.add_heading('Información', level=1)
    doc.add_paragraph(f'Fecha de transcripción: {time.strftime("%Y-%m-%d %H:%M:%S")}')
    if segments:
        doc.add_paragraph(f'Número de segmentos: {len(segments)}')
    
    # Añadir transcripción completa
    doc.add_heading('Transcripción Completa', level=1)
    doc.add_paragraph(text)
    
    # Añadir segmentos con timestamps si están disponibles
    if segments:
        doc.add_heading('Segmentos con Timestamps', level=1)
        for segmento in segments:
            inicio = int(segmento['start'])
            fin = int(segmento['end'])
            texto_seg = segmento['text']
            
            # Formatear tiempo
            inicio_min = inicio // 60
            inicio_sec = inicio % 60
            fin_min = fin // 60
            fin_sec = fin % 60
            
            doc.add_paragraph(f'[{inicio_min:02d}:{inicio_sec:02d} - {fin_min:02d}:{fin_sec:02d}] {texto_seg}')
    
    # Guardar en BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

